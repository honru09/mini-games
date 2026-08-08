"""Apply the visual-commerce P0 closeout to the retained whitepaper DOCX.

This intentionally edits the existing document in place and preserves its
styles, tables, headers, footers and earlier closeout appendix.
"""
from copy import deepcopy
from pathlib import Path
import re
import sys

from docx import Document
from docx.oxml import OxmlElement


DOCX = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(
    r"D:\mini-games\deliverables\Mini Games Platform 产品与技术白皮书 v3.0 完善版.docx"
)
OUTPUT = Path(sys.argv[2]) if len(sys.argv) > 2 else DOCX


def replace_paragraph(paragraph, text):
    p = paragraph._p
    ppr = p.pPr
    first_rpr = None
    for child in list(p):
        if child is ppr:
            continue
        if child.tag.endswith("}r") and child.rPr is not None and first_rpr is None:
            first_rpr = deepcopy(child.rPr)
        p.remove(child)
    run = OxmlElement("w:r")
    if first_rpr is not None:
        run.append(first_rpr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    p.append(run)


def remove_preceding_page_break(paragraph):
    previous = paragraph._p.getprevious()
    if previous is None or not previous.tag.endswith("}p"):
        return False
    visible_text = "".join(previous.itertext()).strip()
    page_breaks = previous.xpath('.//w:br[@w:type="page"]')
    if visible_text or not page_breaks:
        return False
    previous.getparent().remove(previous)
    return True


doc = Document(str(DOCX))
updates = [
    ("v3.0 完善版", "v3.3 视觉商城 P0 收口与 Sticker Cartoon M0 冻结版"),
    ("版本日期：", "版本日期：2026-08-08"),
    (
        "v3.0 当前实现基线 +",
        "v3.3 当前六款/双模式事实基线 + 视觉商城素材 P0 + Sticker Cartoon M0",
    ),
    (
        "Mini Games Platform 已从“五款游戏 Demo”演进为",
        "Mini Games Platform 已从“五款游戏 Demo”演进为具有账号、成长、商城、正式社交图谱和联机房间的网页多人游戏平台。当前产品包含 6 款精选插件化游戏，正式入口只保留人机对战和联机对战；旧同设备多人入口、档案槽位、奖励分支及其对应三语文案已删除。Fast Fun Loop 的约 3 秒选择和约 5 分钟一局仍是体验目标，需以线上冷启动、真实设备和真实网络数据持续校准。",
    ),
    (
        "截至本版，6 款精选游戏人机/联机双模式、",
        "截至本版，6 款精选游戏人机/联机双模式、Seat/Social/Profile v2、游戏外观商城、Daily Task、Replay v1.1、Tournament v1.1、Metrics v2、Reward Resolver、三语与 CI/QA 已具备；六款大厅封面、注册/商城重排、价格契约、五档响应式与本地素材库也已完成自动化及本地浏览器验证。真实 Supabase、真实设备/网络整形、跨实例长期 Metrics、外部 Sentry、远端素材存储和最终 Sticker Cartoon 全量美术仍保持 BLOCKED/待办。",
    ),
    (
        "26. 美术、音频与 Game Feel 的“体验资产化”（v3.0 新增建议）",
        "26. 美术、音频与 Game Feel 体验资产化",
    ),
]

changed = 0
for paragraph in doc.paragraphs:
    current = paragraph.text.strip()
    for prefix, replacement in updates:
        if current.startswith(prefix) and current != replacement:
            replace_paragraph(paragraph, replacement)
            changed += 1
            break

seen_header_parts = set()
for section in doc.sections:
    for header in (section.header, section.first_page_header, section.even_page_header):
        header_key = id(header._element)
        if header_key in seen_header_parts:
            continue
        seen_header_parts.add(header_key)
        for paragraph in header.paragraphs:
            current = paragraph.text.strip()
            normalized = re.sub(r"\s+", " ", current)
            if normalized == "Mini Games Platform · 产品与技术白皮书 v3.0":
                replace_paragraph(paragraph, "Mini Games Platform · 产品与技术白皮书 v3.3")
                changed += 1

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip() == "P3 商业发行":
                replace_paragraph(cell.paragraphs[0], "P3 正式发行")
                changed += 1

for paragraph in list(doc.paragraphs):
    title = paragraph.text.strip()
    if title.startswith((
        "27. 2026-08-08 交接报告执行收口",
        "31. 2026-08-08 视觉商城素材 P0 收口与 Sticker Cartoon M0",
    )) and remove_preceding_page_break(paragraph):
        changed += 1

section_title = "31. 2026-08-08 视觉商城素材 P0 收口与 Sticker Cartoon M0"
if not any(p.text.strip().startswith(section_title) for p in doc.paragraphs):
    doc.add_heading(section_title, level=1)
    doc.add_paragraph(
        "本节以验证提交 52c1a85eedda2651da16b75a8d35a1f8afe3843f、完整 npm test、"
        "五档本地浏览器证据和新《全项目美术风格统一与深度重制执行报告》为准。"
    )
    items = [
        "六款大厅均接入 640×360 / 320×180 WebP、lazy/srcset、完整性哈希、许可与 Emoji fallback；当前六图是可回滚软 3D 过渡批次，不冒充最终 Sticker Cartoon 风格或完整游戏包。",
        "注册与商城完成产品级重排：48 款 Avatar v2（12 免费、36 商城）、单一滚动容器、主预览/试穿、单例弹层、服务端现有价格一致、Starter Background 假购买入口移除和滚动锁回收。",
        "中文、英文、乌克兰语目录统一为 1046 个 key；商品名、游戏顶栏与 Avatar 辅助文本通过连续切换及英文/乌克兰语真实页面泄漏检查。",
        "asset-library 作为 provenance sidecar 记录来源、许可、目录/许可证独立哈希、Prompt/模型、预览与未来对象键；asset_manifest.json 继续是唯一运行时机器事实源。",
        "新目标视觉为 Pocket Tabletop Sticker × Expressive Sticker Cartoon：粗深色轮廓、两级赛璐璐、Q 版强剪影、统一 Facial Kit 与四段式 Motion。只吸收高层视觉语法，不复制商业游戏具体角色、服装、皇冠、构图、表情帧或高潮 Pose。",
        "执行闸门为 Art Bible v1 → Design System v3 / Motion System v1 → Source Manifest v2 → Golden Set（1 Persona×8 状态、4 Avatar、核心 UI、五子棋、飞行棋）→ IP Similarity Review → 批量生产。Golden Set 未通过前不翻新全部 48 Avatar 或其余游戏。",
        "真实 Android、iPhone、Tablet、第二桌面浏览器、真实 Supabase/RLS/并发/备份回滚、真实网络整形和 30 分钟会话尚未执行，因此 Release Candidate 继续 BLOCKED，不得写 production-ready。",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
    changed += len(items) + 2

doc.core_properties.title = "Mini Games Platform 产品与技术白皮书 v3.3 视觉商城 P0 收口与 Sticker Cartoon M0 冻结版"
doc.core_properties.subject = "六款游戏人机/联机事实基线、视觉商城素材 P0 与全项目贴纸卡通重制闸门"
doc.core_properties.comments = "在现有 v3.0 完善版上最小增量修改；验证提交 52c1a85，保留真实环境阻塞项。"
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUTPUT))
print(f"UPDATED {OUTPUT} changes={changed}")
