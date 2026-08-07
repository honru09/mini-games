import sys

from docx import Document


path = r"D:\mini-games\deliverables\Mini Games Platform 产品与技术白皮书 v3.0 完善版.docx"
doc = Document(path)
range_arg = next((arg for arg in sys.argv[1:] if arg.startswith("--range=")), None)
paragraph_range = None
if range_arg:
    start, end = range_arg.split("=", 1)[1].split(":", 1)
    paragraph_range = (int(start), int(end))
terms = [term.lower() for term in sys.argv[1:] if not term.startswith("--range=")]
matches = lambda text: not terms or any(term in text.lower() for term in terms)
print("PARAGRAPHS", len(doc.paragraphs), "TABLES", len(doc.tables), "SECTIONS", len(doc.sections))
for index, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    in_range = paragraph_range is None or paragraph_range[0] <= index <= paragraph_range[1]
    if text and in_range and matches(text):
        print(f"P{index:04d}\t{paragraph.style.name}\t{text}")
for table_index, table in enumerate(doc.tables):
    for row_index, row in enumerate(table.rows):
        cells = " || ".join(cell.text.replace("\n", " / ") for cell in row.cells)
        if matches(cells):
            print(f"T{table_index}R{row_index}\t{cells}")
