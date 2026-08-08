"""Apply small, in-place status corrections to the retained v3.0 whitepaper.

The document is intentionally edited with python-docx instead of regenerated so
its existing page furniture, styles, tables, and numbering remain intact.
"""
from copy import deepcopy
from pathlib import Path
import sys

from docx import Document
from docx.oxml import OxmlElement


DOCX = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(
    r"D:\mini-games\deliverables\Mini Games Platform 产品与技术白皮书 v3.0 完善版.docx"
)
OUTPUT = Path(sys.argv[2]) if len(sys.argv) > 2 else DOCX


def replace_paragraph(paragraph, text):
    """Replace visible text while retaining paragraph style/properties and first-run formatting."""
    p = paragraph._p
    ppr = p.pPr
    first_rpr = None
    for child in list(p):
        if child is ppr:
            continue
        if child.tag.endswith('}r') and child.rPr is not None and first_rpr is None:
            first_rpr = deepcopy(child.rPr)
        p.remove(child)
    run = OxmlElement('w:r')
    if first_rpr is not None:
        run.append(first_rpr)
    node = OxmlElement('w:t')
    if text[:1].isspace() or text[-1:].isspace():
        node.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    node.text = text
    run.append(node)
    p.append(run)


def replace_cell(cell, text):
    if cell.paragraphs:
        replace_paragraph(cell.paragraphs[0], text)
        for extra in cell.paragraphs[1:]:
            replace_paragraph(extra, '')
    else:
        cell.text = text


doc = Document(str(DOCX))

# Prefixes are used because some retained paragraphs contain typographic spaces
# or inline fields that vary between Word/LibreOffice saves.
paragraph_updates = [
    ('Mini Games Platform 已从“五款游戏 Demo”演进为',
     'Mini Games Platform 已从“五款游戏 Demo”演进为具有账号、成长、商城、正式社交图谱和联机房间的网页多人游戏平台。当前产品包含 6 款精选插件化游戏，正式入口只保留人机对战和联机对战；旧同设备多人入口、档案槽位、奖励分支与三语文案已彻底删除。核心体验仍是 Fast Fun Loop：打开约 3 秒进入选择，约 5 分钟完成一局，并能立即再来。'),
    ('截至本版，6 款精选游戏三模式、',
     '截至本版，6 款精选游戏人机/联机双模式、Seat v2、AI Seat/Controller、READY、观战、公开/私密房、房主转移、Social Graph、Profile v2、游戏外观商城、Daily Task、Replay v1.1、Tournament v1.1、Metrics v2、Reward Resolver、三语与 CI/QA 已具备。真实 Supabase、真实设备/网络整形、跨实例长期 Metrics、外部 Sentry 和完整正式美术仍保持 BLOCKED/待办。'),
    ('低门槛复玩。 支持本地热座快速聚会、',
     '低门槛复玩。 支持人机随时开局、AI 进入真实服务端 Seat、联机快速加入与好友邀请；结算后立即重开或在原房间切换游戏。'),
    ('6 款精选游戏均有本地热座、人机 AI、联机三种入口。',
     '6 款精选游戏均有人机 AI 与联机入口；联机房统一使用 human / ai / empty Seat v2，AI 是真实服务端成员而非前端假席位。'),
    ('完整好友/聊天/私聊/举报/公告和投掷表情。',
     '好友申请/接受/忽略/取消/删除、拉黑/解除、固定原因举报、Presence 隐私与 Friends/Online/Recent 社交栏已实现；自由聊天、Feed、公会、处罚/申诉后台仍是后续。'),
    ('3. 大厅先看到在线玩家和开放房间，也可选择本地/AI 直接开局。',
     '3. 大厅先看到在线玩家和开放房间，也可选择人机对战直接开局。'),
    ('单机/AI 结果只允许已认证账号提交自己的结果；',
     'AI 结果只允许已认证账号提交自己的结果；必须带唯一 resultId，并受去重与频控保护。'),
    ('独立观众席、实时快照、重连和赛事编排已上线；用户录像/延迟回放与赛事生产自动建桌仍是后续。',
     '独立观众席、实时快照、重连和自动赛事建桌已上线；Replay v1.1 保存 7 天，支持播放/暂停/跳转/倍速、公开延迟、分享与撤销。高级延迟观战和完整 Tetris 竞技回放仍是后续。'),
    ('独立观众席已完成：中途加入、只读隔离、快照、重连、上限和最终结果；用户延迟回放仍待实现。',
     '独立观众席已完成中途加入、只读隔离、快照、重连、上限和最终结果；Replay v1.1 已实现 7 天记录、播放控制、公开房 5 分钟延迟、分享与撤销。高级延迟观战、防窥屏和完整 Tetris 竞技回放仍待后续。'),
    ('Reward Resolver 按联机/AI/本地三模式分别计算奖励：',
     'Reward Resolver 只按联机与 AI 两种正式模式计算奖励：联机 1v1 胜/平/负为 3/2/1💵 与 12/10/8 XP，多人按名次，AI 为 1/0/0💵 与 8/6/5 XP；赛事积分与普通 💵、XP、胜场隔离。Daily Task 由服务端按 UTC 日期与 claimId 幂等领取。'),
    ('当前 moveLog 只服务断线恢复，且有数量/字节上限；不能直接宣称已完成回放。',
     'Replay v1.1 将正式完成局保存为 7 天版本化记录；仍为有界动作流，moveLog 截断时明确标记，不能将其宣称为完整 Tetris Rule Replay。'),
    ('回放需要版本化事件格式、初始状态、随机种子、参与者公开信息、过期时间和隐私控制。',
     'Replay v1.1 已实现版本、参与者权限、7 天过期、公开房 5 分钟延迟、参与者分享/撤销和仅哈希持久化分享令牌；播放器提供播放/暂停、进度跳转和 0.5–4×。'),
    ('独立观众席已实现只读席位、中途加入、人数上限、初始快照、断线重进和最终结果；延迟策略、观众聊天权限、防窥屏和用户回放仍是竞技模式后续。',
     '独立观众席已实现只读席位、中途加入、人数上限、初始快照、断线重进和最终结果；Replay v1.1 已补齐用户回放和隐私/过期边界。高级观战延迟、聊天权限和防窥屏仍是竞技模式后续。'),
    ('v2.5 当前实现基线 + v3.0 产品 / 体验 / 工程完善规划',
     'v3.0 当前实现基线 + 生产边界、Gameplay Authority 与 AI 持续学习增补'),
    ('9. 社交、事件、回放与观战（完整保留旧规划）',
     '9. 社交、事件、回放与观战（已实现边界 + 后续路线）'),
    ('19. AI 对手 2.0：强度、人格、延迟与成本解耦（v3.0 新增建议）',
     '19. AI 对手 2.0：强度、人格、持续学习、延迟与成本'),
    ('产品范围聚焦为 6 款精选游戏',
     '产品范围聚焦为 6 款精选游戏；井字棋、弹珠跳棋、国际跳棋、斗兽棋、贪吃蛇已从运行时、注册表和 QA 删除。'),
    ('“服务端维护唯一完整游戏状态”更新为真实边界：服务端是房间中继',
     '旧稿“服务端维护唯一完整游戏状态”已按当前协议拆分：普通回合制仍由房间中继与客户端规则协作；Tank 使用 tank-authority-v1 服务端 20Hz 权威模拟，Tetris 使用 tetris-battle-authority-v1 战斗协调权威，象棋棋钟和大富翁拍卖分别由服务端维护。'),
    ('将在线对局拆成休闲社交与未来竞技两个安全等级；当前客户端中继架构继续服务休闲模式',
     '在线对局按能力分层：普通回合制保留轻量中继边界；Tank 正式路径已升级为 tank-authority-v1，Tetris 为 tetris-battle-authority-v1，观众席、赛事、象棋棋钟和大富翁拍卖也已接入服务端协议。旧 relay 仅兼容旧客户端；完整 Tetris 规则重放和更高等级 turn-based 规则验证仍需后续。'),
    ('结算时胜者获得 💵1，每位参与者增加 1 局和 XP',
     '结算由 Reward Resolver v1.0 按模式、名次、有效局资格、首胜/连胜和重复对手规则计算 💵/XP；胜场 wins/totalWins 与余额独立，客户端只展示服务端 Reward Breakdown。'),
    ('可分享回放、观战和赛事直播视图。',
     '独立观众席、实时快照、重连和赛事编排已上线；用户录像/延迟回放与赛事生产自动建桌仍是后续。'),
    ('用户延迟回放仍待实现',
     'Replay v1.1 已实现 7 天记录、播放控制、公开房 5 分钟延迟、分享与撤销；高级延迟观战、防窥屏和完整 Tetris 竞技回放仍待后续。'),
    ('每局胜者 +💵1；平局/失败 +💵0；所有参与者各增加 1 局。',
     'Reward Resolver 只按联机/AI 正式模式计算奖励：联机 1v1 胜/平/负为 3/2/1💵 与 12/10/8 XP，多人按名次，AI 为 1/0/0💵 与 8/6/5 XP；赛事积分与普通 💵、XP、胜场隔离。'),
    ('单机结算服务端不重演真实棋局，因此不能验证胜负真实性',
     '人机对局通过服务端 solo_start 票据、合法动作进度、actionId、resultId 幂等和频控校验；服务端仍不重演普通客户端完整棋局，因此不把客户端 claim 当成完全防作弊。人机决策会在有效赛果后进入个人学习模型。'),
    ('history 是“每位参与者的结算流水”：同一联机 match_id 会按参与者写多行；单机一局写一行。',
     'history 是“每位参与者的结算流水”：同一联机 match_id 会按参与者写多行；人机一局写一行。result_id 用于幂等去重，并非一场联机只允许一行。'),
    ('成长：胜利 +10 XP，未胜 +4 XP；等级由 XP 阈值计算；连胜和最佳连胜由服务端维护。',
     '成长由 Reward Resolver 统一计算：XPNext(level)=min(200,30+5×level)，按模式/名次/资格发放，含首胜、连胜、重复对手衰减和等级里程碑；胜场由 wins/totalWins 独立维护。'),
    ('server/index.js 是零第三方运行依赖的 Node 服务，提供静态文件、POST /api/ai、手写 RFC6455 WebSocket /ws、房间/会话管理、排行榜、结算、商城、JSON 数据回退和可选 Supabase 适配。',
     'server/index.js 是零第三方运行依赖的 Node 服务，提供静态文件、POST /api/ai、手写 RFC6455 WebSocket /ws、房间/会话/结算/商城、Tank/Tetris/观众/赛事/棋钟/拍卖协议、JSON 回退和可选 Supabase 适配；AI 学习与策略知识包拆在 server/ai-learning.js 与 server/ai-strategy-skills.js。'),
    ('每款游戏自己的 scheduleAI() 根据 opts.ai 判断 AI 回合。6 款精选游戏均已接入远端 aiChoose()/DeepSeek 合法选项路径，并保留各自本地算法作为无认证、超时、网络失败或非法返回时的回退。',
     '每款游戏自己的 scheduleAI() 先运行本地强策略，再把合法近优候选交给 aiChoose()/DeepSeek；无认证、超时、网络失败或非法返回时始终回退本地策略，不随机送子。五子棋用威胁空间搜索，象棋用限宽 Alpha-Beta，飞行棋/大富翁用风险与收益评分，坦克用影响图避弹与火线，俄罗斯方块用 Dellacherie 井面和第二块前瞻。'),
    ('服务端代理 DeepSeek，Key 只在环境变量；返回值不在合法选项中时置空并由客户端策略降级。',
     '服务端代理 DeepSeek，Key 只在环境变量；客户端只提交合法选项与归一化特征，返回值不在候选中时置空并由本地策略降级。策略研究依据已写入 server/ai-strategy-skills.js 的维护注释。'),
    ('5 个 persona：傲娇、赌狗、毒舌、萌妹、数学老师；性格影响系统提示与随机性，不改变合法性边界。',
     '5 个 persona：傲娇、赌狗、毒舌、萌妹、数学老师；只影响表达和近优候选中的确定性偏好，不改变合法性、强制胜/防守、服务端权威或奖励。personal-linear-v2 按账号×游戏独立学习，不污染全局 AI。'),
    ('5 个 persona：傲娇、赌狗、毒舌、萌妹、数学老师；只影响表达和近优候选中的确定性偏好',
     '5 个 persona：傲娇、赌狗、毒舌、萌妹、数学老师；只影响表达和近优候选中的确定性偏好，不改变合法性、强制胜/防守、服务端权威或奖励。personal-linear-v2 按账号×游戏隔离；对局中缓存局面哈希、近优候选与归一化特征，赛果后执行胜局强化、败局反事实修正和平局中性反馈。无效局只审计不调权；JSON 与 Supabase ai_learning_models/ai_learning_experiences + apply_ai_learning_v1 支持 resultId 幂等、revision 冲突保护和重启恢复。'),
    ('服务端是 房间中继与结算协调器：客户端各自运行完整游戏逻辑，move 只被校验体积、记录顺序并广播。服务端不模拟 6 款精选游戏的全部合法规则。这是当前产品取舍，不能把客户端签名、token 或双方一致 claim 描述成完全防作弊。',
     '服务端同时承担房间中继与按协议的权威子系统：普通 move 仍由客户端规则层协作；Tank 权威模拟位置/碰撞/伤害/重生/排名，Tetris 权威协调 startAt/seed/目标/垃圾/KO/placement，象棋权威棋钟，大富翁权威拍卖。普通回合制和 Tetris 完整方块规则仍不由服务端全量重演，不能把 token 或一致 claim 描述成完全防作弊。'),
    ('服务端权威内容包括账号会话、房间席位、房主、游戏/人数合法范围、一次性 matchId、有限 moveLog、结果共识、档案、💵、XP、局数、owned、价格和持久化写入。',
     '服务端权威内容包括账号会话、房间席位/房主、matchId、结果共识、档案、💵/XP/胜场、owned/价格、Reward Resolver 与持久化写入；共享协议另权威 Tank 模拟、Tetris 战斗协调、观众只读、赛事状态、象棋棋钟和大富翁拍卖。'),
    ('→ move 中继与有限日志 → 完整结果 claim 共识 → 结算',
     '→ 普通 move 或 Tank/Tetris/棋钟/拍卖专用协议 → 稳定快照/重连 → 完整结果 claim 共识 → Reward Resolver 结算'),
    ('未配置 Supabase 时，服务端使用 data/leaderboard.json；DATA_DIR 可指向测试目录或持久磁盘。',
     '未配置 Supabase 时，服务端使用 DATA_DIR 下的 JSON（含档案、奖励/经济流水和 AI 学习模型）；DATA_DIR 可指向测试目录或持久磁盘。Render 当前未挂载持久盘，不能把 JSON 回退描述为生产持久化。'),
    ('配置 SUPABASE_URL 和仅服务端保存的 service-role secret 后，Node 服务通过 REST 读写 profiles 与 history。',
     '配置 SUPABASE_URL 和仅服务端保存的 service_role secret 后，Node 服务通过 REST 读写 profiles/history/reward_history/economy_ledger/analytics_events，并通过 apply_reward_v1、apply_purchase_v1、apply_ai_learning_v1 原子 RPC 写入奖励、购买和 AI 学习模型/经验。'),
    ('schema 可重复迁移，已对两表启用 RLS',
     'schema 可重复迁移，已对 profiles、history、reward_history、economy_ledger、analytics_events、ai_learning_models、ai_learning_experiences 启用 RLS，并撤销 anon/authenticated 权限；浏览器不得直接读取账号、token 哈希、结算或 AI 学习数据。'),
    ('/api/ai 保留 Origin 白名单、Bearer token、32KB 请求体上限、合法游戏/选项检查、并发与速率限制。',
     '/api/ai 保留 Origin 白名单、Bearer token、32KB 请求体上限、合法游戏/选项/候选特征检查、并发与速率限制；模型只能在本地强策略筛出的近优带内裁决。'),
    ('现有中继不能阻止恶意客户端串通提交一致的伪造走法。若进入竞技/付费赛事，需要服务端规则权威模拟或可验证事件日志。',
     '普通回合制中继仍不能阻止双方串通提交一致的伪造走法；Tank 已具备服务端权威模拟，Tetris 目前是 Battle Coordination Authority，付费/竞技赛事仍需更高等级规则验证和完整回放审计。'),
    ('观战需要只读席位、延迟策略、人数限制、聊天权限和防窥屏',
     '独立观众席已实现只读席位、中途加入、人数上限、初始快照、断线重进和最终结果；Replay v1.1 已提供用户回放、公开延迟和分享权限，高级观战延迟/聊天权限/防窥屏仍是后续。'),
    ('2026-08-07 聚焦执行快照：',
     '2026-08-07 聚焦执行快照：6 款、AI 强度/持续学习、奖励成长、Tank/Tetris/观众/赛事/棋钟/拍卖协议、构建、安全、重连、Supabase adapter、联机 E2E 与 WebSocket 主动断开回归均已通过；真实 Supabase 凭证、完整 Tetris Rule Replay、生产赛事自动建桌和社交/跨端仍待单独验收。'),
    ('专项回归：飞行棋“终点超点折返 + 双端快照一致”；asset manifest、SVG/XML、响应式封面、独立 feature flag、fallback 和六款 runtime ID 一致性；五子棋保持 520×520 逻辑画布与代码网格，俄罗斯方块保持 18×10 二值状态和原快照字段；这些专项不能被通用冒烟替代。',
     '专项回归：飞行棋终点超点折返、六款 AI 状态机/强度、AI 持续学习重启恢复、asset manifest/SVG/XML、Tank Authority、Tetris Battle、观众席、赛事、象棋棋钟、大富翁拍卖、Reward/Schema/Security/Reconnect/E2E；这些专项不能被通用冒烟替代。'),
    ('功能：6 款精选游戏均能开始、操作、结束、重开；本地/AI/联机结果正确。',
     '功能：6 款精选游戏在人机/联机模式均能开始、操作、结束、重开，结果与重连状态正确。'),
    ('建议新增“游客本地试玩”（未来选项）：',
     '未来可评估“游客人机试玩”：只使用程序内 AI fallback，不写服务端成长、不领取 💵、不进入公开社交；当玩家尝试联机、排行榜、商城或跨设备同步时再引导建立 PIN 账号。此方案不恢复同设备多人模式，且必须与当前账号策略兼容并单独评估实现成本。'),
    ('4. 快速开局：AI / 本地 / 创建房间。',
     '4. 快速开局：人机对战 / 快速加入 / 创建房间。'),
    ('5. 实在无人时直接给出本地/AI 快速入口。',
     '5. 实在无人时直接给出人机对战快速入口。'),
    ('新用户打开 → 30 秒内理解平台 → 可本地/AI 立即玩',
     '新用户打开 → 30 秒内理解平台 → 可立即进入人机对战 → 结算 → 创建账号保存成长 → 无空大厅死路。'),
    ('1. 修复并测试飞行棋终点折返，覆盖本地/AI/联机和重连。',
     '1. 修复并测试飞行棋终点折返，覆盖人机/联机和重连。'),
    ('依据近期胜负轻微调整，但必须让用户知道，不用于排行榜竞技。',
     '已实现：personal-linear-v2 按账号×游戏记录局中近优候选；有效胜局强化、败局反事实修正、平局中性反馈，无效局只审计。模型不参与排行榜权威结算，不保存原始完整局面或对话。'),
    ('自适应（未来）：依据近期胜负轻微调整',
     '自适应（已实现个人模型）：依据有效赛果对近优候选做低学习率更新；胜局强化、败局反事实修正、平局中性反馈。用户数据按账号×游戏隔离，不用于排行榜权威结算。'),
    ('当前胜者每局 +💵1 的简单规则适合早期，但随着任务、赛季和活动增加，需要维护经济账本：',
     '旧稿“胜者每局 +💵1”已替换为 Reward Resolver v1.0：按模式、名次、资格、首胜、连胜和重复对手衰减计算 💵/XP，并写入 reward_history/economy_ledger；真实配置集中在 server/reward-engine.js。'),
    ('当前休闲联机仍是“服务端房间中继 + 客户端完整规则状态”。这适合快速迭代，但不应升级宣传为“服务端已验证每一步棋”。',
     '当前联机是分层协议：普通回合制仍为服务端房间中继 + 客户端规则状态；Tank 是 tank-authority-v1，Tetris 是 tetris-battle-authority-v1，象棋/大富翁分别有棋钟/拍卖权威。任何未覆盖的普通规则仍不应宣传为服务端逐步验证。'),
    ('1. 修复飞行棋终点折返 + 双端一致回归。',
     '1. 已完成飞行棋终点折返与双端一致回归，并保留专项测试。'),
    ('观战；',
     '独立观众席已完成：中途加入、只读隔离、快照、重连、上限和最终结果；Replay v1.1 已完成 7 天记录、播放控制、公开延迟与可撤销分享。'),
    ('锦标赛原型；',
     'Tournament v1.1 已完成：六款 3–6 人选择、循环/Swiss、Bye、真实房间自动建桌、积分、重连、自愿弃权和管理员指定判负；赛事积分不写普通经济与胜场。'),
    ('将 Match Event、回放/观战和 Godot 都建立在版本化平台事件协议上，避免形成第二套不兼容逻辑。',
     '观众席、赛事、Tank/Tetris、棋钟和拍卖已使用版本化平台事件协议；用户回放、Match Event 和 Godot 仍按同一协议路线扩展。'),
    ('supabase/schema.sql：profiles/history、幂等字段、RLS 与权限。',
     'supabase/schema.sql：profiles/history/reward/economy/analytics/AI 学习表、奖励/购买/学习原子 RPC、幂等字段、RLS 与权限。'),
]

changed = 0
for paragraph in doc.paragraphs:
    current = paragraph.text.strip()
    for prefix, replacement in paragraph_updates:
        if current.startswith(prefix):
            if current != replacement:
                replace_paragraph(paragraph, replacement)
                changed += 1
            break

table_updates = {
    (0, 1): ['Supabase 持久化', 'schema/RLS、奖励/购买/AI 学习 RPC、服务端适配和状态脚本已备好；真实凭证待提供', '配置 SUPABASE_URL 与 service_role secret，执行迁移、并发/RLS、备份和回滚检查'],
    (0, 2): ['社交系统', '好友申请/接受/忽略/取消/删除、拉黑/解除、举报、Presence 隐私和 Friends/Online/Recent 已完成', '聊天、Feed、公会、处罚/申诉后台'],
    (0, 4): ['观战与回放', '独立观众席 + Replay v1.1：7 天记录、播放器、公开延迟、分享/撤销与权限已完成', '高级观战延迟、防窥屏和完整 Tetris 竞技回放'],
    (1, 1): ['功能完整度', '六款人机/联机、Seat/Social/Profile v2、商城/三语、奖励、每日任务、Replay、赛事和 Metrics 已形成闭环', '聊天/赛季、真实 Supabase、真实设备/网络与跨端仍待办'],
    (1, 3): ['后端架构', '零依赖；Seat v2、权威结算、Replay/Tournament、Metrics v2 与规则 Authority 仍集中在 Node 服务', '继续模块化并完成跨实例持久化/一致性'],
    (1, 5): ['联机逻辑', 'READY、公开/私密房、真人/AI Seat、房主/Controller 转移、观众、赛事弃权/恢复和结果共识已增强', '五子棋/飞行棋完整规则验证与多实例一致性'],
    (1, 6): ['游戏逻辑', '六款规则、AI 强度/持续学习、Game Feel、Tank/Tetris 共享协议和棋钟/拍卖均有专项 QA', '完整 Tetris Rule Replay、更多规则/事件和正式美术'],
    (1, 8): ['可扩展性', '游戏注册表、生命周期、Gameplay Protocol 和资源 manifest 已落地', '用户回放、跨运行时 SDK、社交/跨端继续统一'],
    (2, 0): ['游戏', '人数', 'Seat v2', '人机 AI', '联机', '当前 AI/体验重点'],
    (2, 1): ['五子棋 ⚫', '2', '是', '是', '是', '威胁空间搜索；成五/封堵/双威胁与棋盘层次'],
    (2, 2): ['飞行棋 ✈️', '2–4', '是', '是', '是', '终点/吃子/安全风险；3D 骰子、逐格移动、终点规则'],
    (2, 3): ['迷你大富翁 🏙️', '2–5', '是', '是', '是', '净资产/现金储备/租税风险；掷骰、地块、事件卡、实时拍卖'],
    (2, 4): ['坦克大战 🛡️', '2', '是', '是', '是', '影响图/避弹/火线/BFS 侧翼；服务端实时模拟、射击、爆炸与破坏'],
    (2, 5): ['俄罗斯方块 🧱', '2–4', '是', '是', '是', 'Dellacherie 井面+第二块前瞻；7-Bag、硬降、消行、垃圾与 KO'],
    (2, 6): ['象棋 ♞', '2', '是', '是', '是', '限宽 Alpha-Beta；吃子、将军/将死、合法层与服务端棋钟'],
    (3, 4): ['C→S', 'select_game/start/move + Tank/Tetris/clock/auction 专用消息', '房主控制、普通走子中继和共享 Gameplay Authority 输入'],
    (3, 7): ['S→C', 'peer_status/rejoined/*expired/host_changed + authority snapshots', '重连、超时、房主转移与 Tank/Tetris/棋钟/拍卖权威状态'],
    (4, 4): ['私域', '好友申请/接受/忽略/取消/删除、拉黑/解除、举报和 Presence 隐私已完成', '聊天、Feed、公会与处罚/申诉'],
    (4, 5): ['主页', 'Profile v2：签名、地区、性别标签、在线偏好、精选展示、Avatar/框/背景/Name FX', '徽章墙深度、赛季称号与跨端同步'],
    (5, 1): ['BUG-001', '飞行棋超出终点时不能移动，要求折返', '已修复并转为回归项；人机/联机一致，qa/gameplay-upgrade.js 与 E2E 覆盖', '高', '终点剩余 r、骰子 d>r 时先到终点再后退 d−r；可正常结束'],
    (5, 2): ['BUG-002', '骰子超出时回合不同步', '已修复并转为回归项；完整动画/结算后统一切回合，双端快照和 AI 队列通过', '高', '所有分支完整播放/结算后统一切回合；两端状态、UI、AI 队列一致'],
    (5, 4): ['BUG-004', '持久化配置不当，重启可能丢数据', 'schema/RLS/适配/JSON outbox 已完成；真实凭证、迁移、并发、备份和回滚仍待外部验收', '中', '生产 Supabase 验证通过；无凭证时明确提示；不泄露 service-role secret'],
    (7, 1): ['P0 稳定基线', '数据、货币显示、AI 和高优先级协议', '六款 AI/学习、Reward/独立胜场、Seat/Social、Replay/Tournament/Metrics、规则 Authority、构建、安全与重连回归', '真实 Supabase 和生产协议可审计，仅保留人机/联机入口'],
    (7, 3): ['P1 社交闭环', '从邀请升级到正式关系与安全边界', '好友/拉黑/举报/Presence 已完成；继续聊天、Feed、公会与处罚/申诉', '频控、屏蔽、隐私和三语言持续通过'],
    (7, 4): ['P1 复玩系统', '提升每日回归', 'Daily Task 已完成；继续周任务、赛季与 Match Event', '不破坏公平模式和 5 分钟节奏'],
    (7, 5): ['P1 内容扩展', '增加玩法与观赏性', 'Replay v1.1 已完成；继续高级延迟观战、文字/社交游戏与更多正式美术', '版本化事件和隐私策略保持兼容'],
    (8, 7): ['胜者 +1、所有人 +1 局', '旧规则仅作历史记录；当前由 Reward Resolver 按模式/名次/资格计算 💵、XP 与 wins/totalWins', '已完成并由 reward/security/E2E 回归'],
    (8, 4): ['人机、联机双模式', '旧同设备多人入口、档案槽位、奖励 local 分支与三语文案已删除；DeepSeek persona 保留在 AI 模式', '已完成'],
    (8, 10): ['Supabase + JSON 回退', '保留；schema、奖励/购买/AI 学习 RPC 与 fake adapter 已就绪，真实凭证待提供', '部分完成：需生产迁移、RLS/并发、备份回滚验收'],
    (8, 12): ['好友、邀请、举报', '好友申请/接受/忽略/取消/删除、拉黑/解除、举报、Presence 隐私已完成', '已完成基础闭环；聊天/处罚后台待办'],
    (8, 17): ['回放与观战', 'Replay v1.1 与独立观众席已上线；有界 moveLog 不冒充完整 Tetris Rule Replay', '完成：播放器/隐私/过期/公开延迟/分享；高级竞技回放待办', '已完成 MVP'],
    (8, 21): ['服务端唯一完整状态', '拆分为普通中继 + Tank/Tetris/棋钟/拍卖权威子系统；不宣称所有 turn-based 规则服务端重演', '当前取舍，持续升级规则验证'],
    (9, 2): ['3–15s', '建立身份', '昵称/头像/PIN 一屏完成；解释 PIN 用途', '账号完成后直接进入人机或联机入口'],
    (9, 3): ['15–30s', '选玩法', '默认突出快速加入；显示在线朋友/公开房；无人时提供人机或创建公开房', '无在线玩家时明确提供人机对战或创建公开房'],
    (11, 2): ['飞行棋', '骰子与追逐', '终点超点折返、安全格/碰撞可视化、路径逐格移动已完成', '队伍模式、房规预设、轻事件模式'],
    (11, 4): ['坦克大战', '即时对抗', 'Tank Authority、输入预测/命中反馈、障碍种子、出生保护和移动端摇杆调优', '地图轮换、2v2、观战美术'],
    (11, 5): ['俄罗斯方块', '放置与连击', 'Ghost/Hold、消行/连击、7-Bag、Battle Authority 垃圾/KO 协议已完成', '40L 竞速、完整 Rule Replay、淘汰赛'],
    (11, 6): ['象棋', '深度策略', '将军/将死/被将、合法走法/吃子反馈、序列化和服务端棋钟已完成', '残局、PGN-like 平台棋谱格式、完整规则 Authority'],
    (12, 2): ['策略层', '评估动作', '威胁空间/Alpha-Beta、风险收益、影响图/BFS、Dellacherie 双块前瞻和近优远端裁决'],
    (15, 1): ['L0 当前休闲', '朋友局/低风险', '中继、身份、席位、结果共识、频控；普通回合制仍在此边界'],
    (15, 2): ['L1 可验证日志', '排行榜强化', '事件序列、随机种子、stateHash、双端分叉检测；AI 学习经验只保存哈希/特征'],
    (15, 3): ['L2 规则验证', '重要 turn-based', '后续为普通回合制服务端验证每一步合法性，不负责渲染'],
    (15, 4): ['L3 权威模拟', 'Tank/Tetris/竞技路径', 'Tank 已为服务端状态机/实时 tick；Tetris 为 Battle Coordination，完整规则重放仍待'],
    (17, 3): ['客户端中继被作弊', '中', '高', '异常胜率/结果争议', 'Tank/Tetris/象棋/大富翁已有服务端 Authority；五子棋/飞行棋仍需 L2 规则验证或事件审计'],
    (19, 5): ['俄罗斯方块', '10×18', '7 Tetromino、ghost、消行动效、Battle Authority 垃圾行/KO'],
}

for (table_index, row_index), values in table_updates.items():
    if table_index >= len(doc.tables) or row_index >= len(doc.tables[table_index].rows):
        continue
    row = doc.tables[table_index].rows[row_index]
    for index, value in enumerate(values):
        if index < len(row.cells):
            replace_cell(row.cells[index], value)
            changed += 1

closeout_title = '27. 2026-08-08 交接报告执行收口'
if not any(p.text.strip().startswith(closeout_title) for p in doc.paragraphs):
    doc.add_page_break()
    doc.add_heading(closeout_title, level=1)
    doc.add_paragraph('本节以当前 main 源码、正常用户入口和本轮专项/全量 QA 为准，替代旧稿中已过期的同设备多人、Replay 未实现、Metrics 仅 API 和赛事无恢复入口等描述。')
    closeout_items = [
        '玩法收口：六款游戏只保留人机与联机入口；旧同设备多人档案槽位、奖励 local 分支、三语词条、CSS/HTML 和本地结算测试已删除。',
        'Seat / Social / Profile：human / ai / empty Seat v2、READY、AI Controller、真人接管、房主转移、公开/私密房、观战、好友/拉黑/举报、Presence 隐私与 Avatar/背景/收藏试穿已进入主线。',
        '产品闭环：六款游戏外观商城、Daily Task 服务端幂等领取、Replay v1.1 分享/延迟/过期、Tournament v1.1 自愿弃权/管理员指定判负、Metrics v2 管理页面/历史/CSV/阈值/脱敏错误闭环均已实现。',
        '安全边界：分享令牌只持久化哈希；赛事积分不进入普通 💵、XP、胜场；Metrics 全路由 Bearer 鉴权、限频并只记录聚合值、错误类别和不可逆 IP 摘要。',
        '验证：i18n、DOM、Reward、Metrics、Replay、Tournament Recovery 专项和完整 npm test 均需通过后方可接受；public/index.html 只能由 build 生成。',
        '仍未完成：真实 Supabase/RLS/并发/备份回滚、真实移动设备/第二桌面浏览器、真实网络整形、30 分钟 Synthetic Session、跨实例长期 Metrics、外部 Sentry、完整正式美术与跨端发行。Release Candidate 因这些真实环境闸门继续保持 BLOCKED。',
    ]
    for item in closeout_items:
        doc.add_paragraph(item, style='List Bullet')
    changed += len(closeout_items) + 2

doc.core_properties.title = 'Mini Games Platform 产品与技术白皮书 v3.2 交接执行收口版'
doc.core_properties.subject = '六款游戏人机/联机平台；Seat/Social/Profile、Reward、Replay、Tournament 与 Metrics 事实基线'
doc.core_properties.comments = '按当前源码、专项/全量 QA 与真实阻塞项更新；保留原版布局与样式。'
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUTPUT))
print(f'UPDATED {OUTPUT} changes={changed}')
